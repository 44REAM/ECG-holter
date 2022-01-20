
import os
import shutil

import pytz
from docx.shared import Pt
import matplotlib.pyplot as plt
from matplotlib.dates import date2num, AutoDateLocator
import matplotlib.dates as mdates
from docx.shared import Inches, Cm
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from scipy.ndimage import zoom

from .utils import *
from .. import peakobject as po

from docx.shared import Pt

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def make_rows_center(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

def plot_datetime(savepath, y_signals, x_time, title, ylabel, tz, ymin, ymax, color = 'b'):

    if len(x_time)<=1:
        return

    plt.close()
    fig, ax = plt.subplots(figsize=(30, 8))
    ax.set_xlim(x_time[0], x_time[-1])
    ax.plot(date2num(x_time), y_signals, color)

    plt.title(title, fontsize = 24)
    ax.set_xlim(x_time[0],x_time[-1])
    ax.set_ylim(ymin,ymax)
    ax.set_ylabel(ylabel, fontsize=22)

    # auto format date
    # ax.xaxis.set_major_formatter( AutoDateFormatter(locator) )
    locator = AutoDateLocator()
    ax.xaxis.set_major_locator(locator)

    # manual format date
    date_format  = mdates.DateFormatter('%d %H:%M:%S', tz = tz)
    ax.xaxis.set_major_formatter(date_format)
    ax.tick_params(axis='both', which='major', labelsize=16)

    ax.grid(True)
    plt.savefig(savepath, bbox_inches='tight')

min_max = 3
y_tick_major = np.arange(-4, 4,0.5)
y_tick_minor = np.arange(-4, 4,0.1)

def plot_ecg_10_sec(savepath, y_signals_ch2, y_signals_ch1, x_time, interval_array, interval_time_array, annotation, tz):
    
    if len(x_time)<=1:
        return
    
    plt.close()
    date_format  = mdates.DateFormatter('%d %H:%M:%S', tz = tz)

    fig, ax = plt.subplots(2, figsize=(30, 15))

    # ax.set_ylim([-min_max, min_max])
    for a in ax:

        a.set_xticks(x_time[:: 100])
        a.set_xticks(x_time[:: 20], minor=True)
        a.set_yticks(y_tick_major)
        a.set_yticks(y_tick_minor, minor=True)
        a.grid(which='major')
        a.grid(which='minor', alpha=0.3)

        a.xaxis.set_major_formatter(date_format)
        a.tick_params(axis='both', which='major', labelsize=16)
        a.set_ylim(-min_max, min_max)
        a.set_xlim(x_time[0], x_time[-1])
        a.set_ylabel("mV", fontsize=22)

        for idx,t in enumerate(a.xaxis.get_major_ticks()):
            if idx%5==0:
                continue
            t.label.set_visible(False)
            t.tick1On = False # <----

        for interval, time_interval, anno in zip(interval_array[1:], interval_time_array[1:], annotation[1:]):

            xmax = timestamp_msec_to_datetime(time_interval)
            xcenter = timestamp_msec_to_datetime(time_interval-interval/2)
            a.vlines(x = xmax, ymin = 2, ymax = 2.5 ,color = "#f70000")
            a.text(xcenter , 2,  interval, fontsize = 16, ha='center')
            a.text(xmax , 2.5,  anno, fontsize = 16, ha='center')

    ax[0].set_title("CHANNEL I", fontsize = 24)
    ax[1].set_title("CHANNEL II", fontsize = 24)

    ax[0].plot(x_time, y_signals_ch1)
    ax[1].plot(date2num(x_time), y_signals_ch2)
    plt.savefig(savepath, bbox_inches='tight')

def plot_ecg_1_hour(savepath, y_signals, x_time, annotation_time, title, tz):
    plt.close()
    y_scatter = np.ones(len(annotation_time))*(-4)

    fig, ax = plt.subplots(figsize=(30, 8))
    ax.set_xlim(x_time[0], x_time[-1])
    ax.plot(date2num(x_time), y_signals)
    ax.scatter(annotation_time, y_scatter, c = "#f70000", zorder = 1)

    plt.title(title, fontsize = 24)
    ax.set_xlim(x_time[0],x_time[-1])
    ax.set_ylim(-5,5)
    ax.set_ylabel("mV", fontsize=22)

    # auto format date
    # ax.xaxis.set_major_formatter( AutoDateFormatter(locator) )
    locator = AutoDateLocator()
    ax.xaxis.set_major_locator(locator)

    # manual format date
    date_format  = mdates.DateFormatter('%d %H:%M:%S', tz = tz)
    ax.xaxis.set_major_formatter(date_format)
    ax.tick_params(axis='both', which='major', labelsize=16)

    ax.grid(True)
    plt.savefig(savepath, bbox_inches='tight')


def report_diagnosis1hour(savefolder,signals, peakarray: po.PeakArray, diag ,segment, tz):
    
    # Plot HR 24 hour
 
    fs = peakarray.fs
    n_segment = len(segment)
    start = 0

    for i in range(25):
        if i>=n_segment:
            break
        landmark_peak = peakarray[start:segment[i]]
        diag_location = landmark_peak.get_diagnosis_time(diag)

        if len(diag_location) == 0:
            continue

        landmark_signal = signals[landmark_peak[0].idx: landmark_peak[-1].idx]

        time_array = sample_to_sec(np.arange(landmark_peak[0].idx, landmark_peak[-1].idx), fs) + peakarray.START_TS/1000
        time_array = [timestamp_sec_to_datetime(ts) for ts in time_array]

        savepath = os.path.join(savefolder, f'{po.ANNOTATION[diag]}_hour_{(i+1):02d}.jpg')
        plot_ecg_1_hour(savepath, landmark_signal, time_array,diag_location ,f"{po.ANNOTATION[diag]} Hour {(i+1):02d}", tz)


        start = segment[i]

def report_diagnosis(savefolder, peakarray: po.PeakArray, signals, signals_ch1, diag: po.ECGLabel, tz, limit = 10000):
    segment = peakarray.get_segment_index_every_nsec(nsec = 10)
    start = 0
    fs=peakarray.fs
    picture_count = 0
    for end in segment:
        if picture_count >= limit:
            break
        landmark_peak = peakarray[start:end]
        start = end

        out = False
        have_diag = False
        annotation = []
        for peak in landmark_peak:

            if not isinstance(peak, po.QRS):
                out = True
                break

            if peak.diagnosis == po.ECGLabel.UNKNOWN:
                out = True
                break

            if peak.diagnosis == diag:
                have_diag = True

            annotation.append(peak.get_annotation())

        if out or (not have_diag):
            continue

        r_peak = landmark_peak.get_r_index()
        if len(r_peak)<4:
            continue

        start_idx = landmark_peak[0].idx
        end_idx = landmark_peak[-1].idx

        length = end_idx - start_idx
        add = sec_to_sample(10, fs) -length

        end_idx +=add+1

        landmark_signals = signals[start_idx: end_idx]
        landmark_signals_ch1 = signals_ch1[start_idx: end_idx]

        time_array = sample_to_sec(np.arange(start_idx, end_idx), fs) + peakarray.START_TS/1000
        time_array = [timestamp_sec_to_datetime(ts) for ts in time_array]
        interval, time_interval = landmark_peak.get_interval(to_time=True, report=True)

        savepath = os.path.join(savefolder, f"{time_array[0].strftime('%H.%M.%S (%d_%m_%Y)')}.jpg")
        plot_ecg_10_sec(savepath, landmark_signals, landmark_signals_ch1, time_array, interval, time_interval, annotation, tz)
        picture_count+=1

def report_hr(savefolder, peakarray: po.PeakArray, segment, tz):
    
    # Plot HR 24 hour
    peak_hr, peak_time = peakarray.get_hr_with_time()

    savepath = os.path.join(savefolder, f'0_hr_total.jpg')
    plot_datetime(savepath, peak_hr, peak_time, f"Total HR plot", "HR", tz, 0, 200)

    n_segment = len(segment)
    start = 0
    for i in range(25):
        if i>=n_segment:
            break

        peak_hr, peak_time = peakarray.get_hr_with_time(start_idx=start, end_idx=segment[i])
        savepath = os.path.join(savefolder, f'hr{(i+1):02d}.jpg')
        plot_datetime(savepath, peak_hr, peak_time, f"HR plot one hour {(i+1):02d}", "HR", tz, 0,200)
        start = segment[i]

def report_question(savefolder, peakarray: po.PeakArray, n_signal, tz):

    mask = get_section_mask(peakarray.Q_SECTION, n_signal, invert =True)

    ts_array = sample_to_sec(np.arange(n_signal), peakarray.fs)  + peakarray.START_TS/1000
    ts_array = [timestamp_sec_to_datetime(ts) for ts in ts_array]
    savepath = os.path.join(savefolder, f"0_Total")
    plot_datetime(savepath, mask,ts_array, "Total signal lost",'one = lost', tz, 0, 1.2, color ='r' )

def clean_folder(folder_path):
    print(folder_path)
    if os.path.exists(folder_path):
        shutil.rmtree(folder_path)
        os.mkdir(folder_path)
    else:
        os.mkdir(folder_path)


def format_diagnosis(document, savefolder,folder_name,
                 peakarray,signals, signals_ch1, 
                 signals_resample, peakarray_resample, diag,
                 segment, tz, document_header, save_pic_limit = 5, document_n_pic = 3,
                 document_pic_width = 7):

    document.add_page_break()
    document.add_heading(document_header, level = 1)
    document.add_paragraph()
    folder_path = os.path.join(savefolder, folder_name)
    clean_folder(folder_path)
    hour_folder_path = os.path.join(folder_path, "hour_report")
    clean_folder(hour_folder_path)
    report_diagnosis1hour(hour_folder_path, signals_resample, peakarray_resample, diag, segment, tz)
    report_diagnosis(folder_path, peakarray, signals, signals_ch1, diag, tz, limit = save_pic_limit)

    for picture in os.listdir(hour_folder_path):
  
        pic_path = os.path.join(hour_folder_path, picture)
        document.add_picture(pic_path, width=Inches(document_pic_width))

    document.add_heading("Example", level = 5)

    for picture in os.listdir(folder_path)[:document_n_pic]:
        pic_path = os.path.join(folder_path, picture)
        if os.path.isdir(pic_path):
            continue

        document.add_picture(pic_path, width=Inches(document_pic_width))
    if len(os.listdir(folder_path)) == 0:
        document.add_paragraph("There are no this labelpeak or consider a noise")

def report_document(savefolder, peakarray: po.PeakArray, signals, signals_ch1, report_lost_signal = False):

    n_signals = len(signals)
    n_peak = len(peakarray.get_r_index())

    signals_resample = zoom(signals, 10/peakarray.fs)
    peakarray_resample = po.resample(peakarray, 10)



    if peakarray.START_TS is None:
        raise ValueError("Start time not yet assigned")

    if peakarray.Q_SECTION is None:
        raise ValueError("Question not yet assigned")

    # get timezone
    tz = pytz.timezone('Asia/Bangkok')
    # get segment every 1 hour
    segment = peakarray.get_segment_index_every_nsec(nsec = 3600)

    # setting folder
    # -------- Report HR folder
    hr_folder = "hr_report"
    hr_folder_path = os.path.join(savefolder, hr_folder)

    clean_folder(hr_folder_path)
    report_hr(hr_folder_path, peakarray, segment, tz)

    # --------Report Diag
    
    

    # Create Document
    document = Document()

    document.add_heading("ECG report", level=0)
    document.add_heading("Information", level = 2)
    document.add_paragraph()

    start_time = timestamp_msec_to_datetime(peakarray.START_TS, return_object=False)
    document.add_paragraph(f"Start time: {start_time}")
    time_signal = sec_to_time_format(sample_to_sec(n_signals, peakarray.fs))
    document.add_paragraph(f"Duration: {time_signal}")
    document.add_paragraph(f"Total signal: {n_signals} (Sampling rate {peakarray.fs}Hz)")
    n_lost_signal = count_section_sample(peakarray.Q_SECTION)
    document.add_paragraph(f"Lost signal: {n_lost_signal} ({round(n_lost_signal/n_signals*100, 2)}%)")


    document.add_paragraph()
    document.add_heading("General summary", level = 2)
    document.add_paragraph()
    document.add_paragraph(f"Total peak: {n_peak}")

    table = document.add_table(rows=11, cols=2, style = "Table Grid")
    # Heading
    table.cell(0, 0).text = "Label"
    table.cell(0, 1).text = "Count"
    make_rows_bold(table.rows[0])
    make_rows_center(table.rows[0])

    # Diagnosis
    all_labels = peakarray.get_all_labels()
    table.cell(1, 0).text = "Normal"
    table.cell(2, 0).text = "Tachycardia"
    table.cell(3, 0).text = "Bradycardia"
    table.cell(4, 0).text = "Pause"
    table.cell(5, 0).text = "PVC"
    table.cell(6, 0).text = "PAC"
    table.cell(7, 0).text = "AF"
    table.cell(8, 0).text = "VTVF"
    table.cell(9, 0).text = "Unknown"
    table.cell(10, 0).text = "Question"

    table.cell(1, 1).text = str(len(all_labels['normal']))
    table.cell(2, 1).text = str(len(all_labels['tachy']))
    table.cell(3, 1).text = str(len(all_labels['brady']))
    table.cell(4, 1).text = str(len(all_labels['pause']))
    table.cell(5, 1).text = str(len(all_labels['pvc']))
    table.cell(6, 1).text = str(len(all_labels['pac']))
    table.cell(7, 1).text = str(len(all_labels['af']))
    table.cell(8, 1).text = str(len(all_labels['ventricular']))
    table.cell(9, 1).text = str(len(all_labels['unknown']))
    table.cell(10, 1).text =str(len(all_labels['question']))

    # Metric
    document.add_paragraph()
    table = document.add_table(rows=9, cols=2, style = "Table Grid")

    all_metrics = peakarray.get_all_metrics( rm_outlier=True, upper_percentile=95, lower_percentile = 5)
    hrmax_time = timestamp_msec_to_datetime(all_metrics['hrmax_ts'], mode = 'bkk', return_object=False)
    hrmin_time = timestamp_msec_to_datetime(all_metrics['hrmin_ts'], mode = 'bkk', return_object=False)


    # Diagnosis
    table.cell(0, 0).text = "Mean HR"
    table.cell(0, 1).text = str(all_metrics['hr'])
    table.cell(1, 0).text = "Max HR"
    table.cell(1, 1).text = f"{all_metrics['hrmax']}     at {hrmax_time}" 
    table.cell(2, 0).text = "Min HR"
    table.cell(2, 1).text = f"{all_metrics['hrmin']}     at {hrmin_time}" 
    table.cell(3, 0).text = "SDNN"
    table.cell(3, 1).text = str(all_metrics['sdnn'])
    table.cell(4, 0).text = "ASDNN"
    table.cell(4, 1).text = str(all_metrics['asdnn'])
    table.cell(5, 0).text = "SDANN"
    table.cell(5, 1).text = str(all_metrics['sdann'])
    table.cell(6, 0).text = "RMSSD"
    table.cell(6, 1).text = str(all_metrics['rmssd'])
    table.cell(7, 0).text = "QT"
    table.cell(7, 1).text = str(all_metrics['qt'])
    table.cell(8, 0).text = "QTc"
    table.cell(8, 1).text = str(all_metrics['qtc'])
    make_rows_bold(table.columns[0])

    pic_width = 7

    # HR picture
    document.add_page_break()
    document.add_heading("HR report", level = 1)
    document.add_paragraph()
    for picture in os.listdir(hr_folder_path):
        pic_path = os.path.join(hr_folder_path, picture)
        document.add_picture(pic_path, width=Inches(pic_width))

    n_pic = 4


    save_pic_limit = 100000
    save_pic_limit_normal = 20
    segment = peakarray_resample.get_segment_index_every_nsec(nsec = 3600)

    document.add_page_break()
    document.add_heading("Normal report", level = 1)
    document.add_paragraph()
    folder_path = os.path.join(savefolder, "normal_report")
    clean_folder(folder_path)
    report_diagnosis(folder_path, peakarray, signals, signals_ch1, po.ECGLabel.NORMAL, tz, limit = save_pic_limit_normal)
    
    for picture in os.listdir(folder_path)[:n_pic]:
        pic_path = os.path.join(folder_path, picture)
        document.add_picture(pic_path, width=Inches(pic_width))
    if len(os.listdir(folder_path)) == 0:
        document.add_paragraph("There are no this labelpeak or consider a noise")

    format_diagnosis(document, savefolder, "brady_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.BRADYCARDIA, segment, tz, "Bradycardia report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)


    format_diagnosis(document, savefolder, "tachy_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.TACHYCARDIA, segment, tz, "Tachycardia report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)

    format_diagnosis(document, savefolder, "pause_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.PAUSE, segment, tz, "Pause report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)


    # ------------------------------------------------------------------------------
    format_diagnosis(document, savefolder, "pvc_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.PVC, segment, tz, "PVC report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)

    # ------------------------------------------------------------------------------
    format_diagnosis(document, savefolder, "pac_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.PAC, segment, tz, "PAC report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)

    format_diagnosis(document, savefolder, "af_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.AF, segment, tz, "AF report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)

    format_diagnosis(document, savefolder, "vtvf_report",
            peakarray,signals,signals_ch1, signals_resample, 
            peakarray_resample,po.ECGLabel.VENTRICULAR, segment, tz, "VTVF report",
            save_pic_limit = save_pic_limit, document_n_pic=n_pic, document_pic_width=pic_width)


    # Question picture
    if report_lost_signal:
        lost_folder = "question_report"
        lost_folder_path = os.path.join(savefolder, lost_folder)

        if os.path.exists(lost_folder_path):
            shutil.rmtree(lost_folder_path)
            os.mkdir(lost_folder_path)
        else:
            os.mkdir(lost_folder_path)
        report_question(lost_folder_path, peakarray, n_signals, tz)
        document.add_page_break()
        document.add_heading("Lost signal report", level = 1)
        document.add_paragraph()
        for picture in os.listdir(lost_folder_path):
            pic_path = os.path.join(lost_folder_path, picture)
            document.add_picture(pic_path, width=Inches(7))


    for section in document.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


    document.save(os.path.join(savefolder, "report.docx"))